"""
DocumentIngestion — multi-format document ingestion with metadata.

Supports: PDF, TXT, MD, DOCX
"""

import os
from datetime import datetime, timezone

from sqlalchemy.orm import Session

from app.db.session import SessionLocal
from app.models.document import Document
from app.rag.embedder import Embedder
from app.rag.pdf_loader import PDFLoader
from app.rag.text_splitter import TextSplitter


def _load_text_file(file_path: str) -> str:
    """Load plain text / markdown files."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _load_docx_file(file_path: str) -> str:
    """Load DOCX files using python-docx."""
    try:
        import docx
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX support. "
            "Install it with: pip install python-docx"
        )


def load_document(file_path: str, file_type: str) -> str:
    """Dispatch to the correct loader based on file type."""
    if file_type == "pdf":
        return PDFLoader.load(file_path)
    elif file_type in ("txt", "md"):
        return _load_text_file(file_path)
    elif file_type == "docx":
        return _load_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


class DocumentIngestion:

    def __init__(self):
        self.embedder = Embedder()
        self.splitter = TextSplitter()

    def ingest(self, file_path: str) -> dict:
        """
        Ingest any supported document into the vector store.

        Returns a summary dict with filename, file_type, and chunk_count.
        """
        filename = os.path.basename(file_path)
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

        # map extensions to canonical type
        ext_map = {"pdf": "pdf", "txt": "txt", "md": "md", "docx": "docx"}
        file_type = ext_map.get(ext, "txt")

        text = load_document(file_path, file_type)
        chunks = self.splitter.split_text(text)
        total_chunks = len(chunks)

        db: Session = SessionLocal()
        try:
            for idx, chunk in enumerate(chunks):
                embedding = self.embedder.embed(chunk)
                doc = Document(
                    content=chunk,
                    embedding=embedding,
                    filename=filename,
                    file_type=file_type,
                    chunk_index=idx,
                    total_chunks=total_chunks,
                    uploaded_at=datetime.now(timezone.utc),
                )
                db.add(doc)
            db.commit()
        finally:
            db.close()

        return {
            "filename": filename,
            "file_type": file_type,
            "chunks": total_chunks,
        }

    # backward-compat alias used by older code
    def ingest_pdf(self, file_path: str) -> int:
        result = self.ingest(file_path)
        return result["chunks"]
